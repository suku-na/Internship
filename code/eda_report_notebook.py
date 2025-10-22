
"""
EDA report generator (eda_report_notebook.py)
With tabular numeric summary and compact plots.
"""

import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from io import BytesIO
from docx import Document
from docx.shared import Inches
import numpy as np

sns.set(style="whitegrid")

def generate_plot_description(df, plot_type: str) -> str:
    num_cols = df.select_dtypes(include=np.number).columns.tolist()
    cat_cols = df.select_dtypes(exclude=np.number).columns.tolist()
    if plot_type == "hist":
        if not num_cols: return "No numeric columns."
        col = num_cols[0]; s = df[col].dropna()
        if s.empty: return f"No data for {col}."
        skew = s.skew()
        skew_text = "symmetric" if abs(skew)<0.5 else "right-skewed" if skew>0.5 else "left-skewed"
        return f"Distribution of **{col}** is {skew_text}, mean={s.mean():.2f}, median={s.median():.2f}."
    if plot_type == "box":
        if not num_cols: return "No numeric columns."
        col = num_cols[0]; s=df[col].dropna()
        if s.empty: return f"No data for {col}."
        q1,q3=np.percentile(s,[25,75]); iqr=q3-q1
        out=((s<q1-1.5*iqr)|(s>q3+1.5*iqr)).sum()
        return f"Boxplot of **{col}**: IQR={iqr:.2f}, ~{int(out)} outliers."
    if plot_type=="heatmap":
        if len(num_cols)<2: return "Not enough numeric cols."
        corr=df[num_cols].corr()
        vals=corr.where(np.triu(np.ones(corr.shape),k=1).astype(bool)).unstack().dropna().abs().sort_values(ascending=False)
        if not vals.empty:
            pair=vals.index[0]; val=vals.iloc[0]
            return f"Strongest correlation: **{pair[0]}** vs **{pair[1]}** (r={val:.2f})."
        return "Correlation heatmap summary."
    if plot_type=="bar":
        if not cat_cols: return "No categorical columns."
        col=cat_cols[0]; vc=df[col].value_counts().head(3)
        return f"Top in **{col}**: "+", ".join([f"{idx} ({val})" for idx,val in vc.items()])
    return "Plot description not available."

def _fig_to_bytes(fig):
    buf=BytesIO(); fig.savefig(buf,format="png",bbox_inches="tight"); plt.close(fig); buf.seek(0); return buf

def analyze_and_report(df: pd.DataFrame, topn:int=10, output_path:str=None)->bytes:
    doc=Document()
    doc.add_heading("Exploratory Data Analysis Report",0)
    # Executive summary
    doc.add_heading("Executive Summary",level=1)
    doc.add_paragraph(f"- Dataset has {df.shape[0]} rows and {df.shape[1]} columns.")
    miss=df.isna().sum()
    if (miss>0).any():
        col=miss.idxmax()
        doc.add_paragraph(f"- Column {col} has most missing: {miss[col]} ({miss[col]/len(df)*100:.1f}%).")
    num=df.select_dtypes(include=np.number)
    if not num.empty:
        doc.add_paragraph(f"- Numeric column with max variance: {num.var().idxmax()}.")
    cat=df.select_dtypes(exclude=np.number)
    if not cat.empty:
        col=cat.columns[0]; doc.add_paragraph(f"- Most frequent in {col}: {df[col].value_counts().idxmax()}.")
    # Dataset overview
    doc.add_heading("Dataset Overview",level=1)
    types=df.dtypes.astype(str).to_frame("dtype"); types["missing"]=df.isna().sum()
    table=doc.add_table(rows=1,cols=3); hdr=table.rows[0].cells; hdr[0].text="Column"; hdr[1].text="Type"; hdr[2].text="Missing"
    for c,r in types.iterrows():
        row=table.add_row().cells; row[0].text=str(c); row[1].text=r['dtype']; row[2].text=str(r['missing'])
    # Numeric summary as table
    if not num.empty:
        doc.add_heading("Numeric Summary",level=1)
        stats=num.describe().T.round(2).reset_index()
        table=doc.add_table(rows=1,cols=len(stats.columns)); table.style="Light Grid"
        for i,col in enumerate(stats.columns): table.rows[0].cells[i].text=col
        for _,r in stats.iterrows():
            row=table.add_row().cells
            for i,val in enumerate(r): row[i].text=str(val)
    # Categorical summary
    if not cat.empty:
        doc.add_heading("Categorical Summary",level=1)
        for c in cat.columns[:topn]:
            vc=df[c].value_counts().head(5)
            doc.add_paragraph(f"{c}",style="Heading 2")
            table=doc.add_table(rows=1,cols=2); table.style="Light Grid"; table.rows[0].cells[0].text="Value"; table.rows[0].cells[1].text="Count"
            for idx,val in vc.items():
                row=table.add_row().cells; row[0].text=str(idx); row[1].text=str(val)
    # Correlation heatmap
    if num.shape[1]>=2:
        corr=num.corr(); fig,ax=plt.subplots(figsize=(4.5,3))
        sns.heatmap(corr,ax=ax,annot=True,fmt=".2f",cmap="RdBu_r",center=0)
        doc.add_heading("Correlation Heatmap",level=1)
        doc.add_paragraph(generate_plot_description(df[num.columns],"heatmap"))
        doc.add_picture(_fig_to_bytes(fig),width=Inches(4.5))
    # Distributions
    if not num.empty:
        doc.add_heading("Distributions",level=1)
        for col in num.columns[:6]:
            fig,ax=plt.subplots(figsize=(3.5,2.5)); sns.histplot(df[col].dropna(),kde=True,ax=ax)
            ax.set_title(f"Distribution of {col}")
            doc.add_picture(_fig_to_bytes(fig),width=Inches(3.5))
            doc.add_paragraph(generate_plot_description(df[[col]],"hist"))
    # Boxplots
    if not num.empty:
        doc.add_heading("Boxplots",level=1)
        for col in num.columns[:6]:
            fig,ax=plt.subplots(figsize=(3.5,2)); sns.boxplot(x=df[col],ax=ax)
            ax.set_title(f"Boxplot of {col}")
            doc.add_picture(_fig_to_bytes(fig),width=Inches(3.5))
            doc.add_paragraph(generate_plot_description(df[[col]],"box"))
    # Categorical charts
    if not cat.empty:
        doc.add_heading("Top Categories",level=1)
        for c in cat.columns[:topn]:
            vc=df[c].value_counts().head(10)
            if vc.empty: continue
            fig,ax=plt.subplots(figsize=(3.5,2.5)); sns.barplot(x=vc.values,y=[str(x) for x in vc.index],ax=ax)
            ax.set_title(f"Top categories in {c}")
            doc.add_picture(_fig_to_bytes(fig),width=Inches(3.5))
            doc.add_paragraph(generate_plot_description(df[[c]],"bar"))
    # Conclusion
    doc.add_heading("Conclusion",level=1)
    doc.add_paragraph("This report summarizes dataset structure, missing values, distributions, correlations, and categories. Further analysis (feature engineering, statistical tests, modeling) is recommended for deeper insights.")
    bio=BytesIO(); doc.save(bio); bio.seek(0)
    if output_path:
        with open(output_path,"wb") as f: f.write(bio.getvalue())
    return bio.getvalue()
